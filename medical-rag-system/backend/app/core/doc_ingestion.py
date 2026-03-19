"""文档导入：类型推断、文本提取与原始文件持久化。"""

from __future__ import annotations

import io
import hashlib
import mimetypes
import shutil
import uuid
from datetime import date, datetime
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from app.core.config import settings

from .chunker import CHUNK_STRATEGY_VERSION
from .text_utils import clean_text, extract_title, now_iso, strip_html


@dataclass
class ExtractedDocument:
    text: str
    viewer_mode: str
    mime_type: str
    meta: Dict[str, Any]


class DuplicateDocumentError(ValueError):
    def __init__(self, existing_doc: Dict[str, Any]):
        self.existing_doc = existing_doc
        title = existing_doc.get("title") or existing_doc.get("original_file_name") or existing_doc.get("doc_id") or "该文档"
        super().__init__(f"知识库已有《{title}》，请确认是否覆盖。")


def guess_doc_type(file_name: str) -> str:
    lower = file_name.lower()
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".html") or lower.endswith(".htm"):
        return "html"
    if lower.endswith(".docx"):
        return "docx"
    return "text"


def guess_mime_type(file_name: str, doc_type: str) -> str:
    guessed, _ = mimetypes.guess_type(file_name)
    if guessed:
        return guessed
    if doc_type == "pdf":
        return "application/pdf"
    if doc_type == "docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if doc_type == "html":
        return "text/html; charset=utf-8"
    return "text/plain; charset=utf-8"


def _join_segments(segments: List[Tuple[int, str]], key_name: str) -> Tuple[str, List[Dict[str, int]]]:
    parts: List[str] = []
    spans: List[Dict[str, int]] = []
    cursor = 0
    total = len(segments)

    for idx, (label, text) in enumerate(segments):
        parts.append(text)
        start = cursor
        end = start + len(text)
        spans.append({key_name: label, "start": start, "end": end})
        cursor = end
        if idx < total - 1:
            parts.append("\n\n")
            cursor += 2

    return "".join(parts), spans


def _extract_pdf(content: bytes, file_name: str) -> ExtractedDocument:
    try:
        import fitz  # type: ignore

        pages: List[Tuple[int, str]] = []
        with fitz.open(stream=content, filetype="pdf") as pdf:
            for page_number, page in enumerate(pdf, start=1):
                page_text = clean_text(page.get_text("text"))
                pages.append((page_number, page_text))

        text, page_spans = _join_segments(pages, "page")
        return ExtractedDocument(
            text=text,
            viewer_mode="pdf",
            mime_type=guess_mime_type(file_name, "pdf"),
            meta={"page_spans": page_spans},
        )
    except Exception:
        fallback_text = content.decode("utf-8", errors="ignore")
        return ExtractedDocument(
            text=clean_text(fallback_text),
            viewer_mode="structured_text",
            mime_type=guess_mime_type(file_name, "pdf"),
            meta={},
        )


def _extract_docx(content: bytes, file_name: str) -> ExtractedDocument:
    try:
        from docx import Document  # type: ignore

        doc = Document(io.BytesIO(content))
        paragraphs: List[Tuple[int, str]] = []
        para_index = 0
        for para in doc.paragraphs:
            cleaned = clean_text(para.text)
            if not cleaned:
                continue
            para_index += 1
            paragraphs.append((para_index, cleaned))

        text, paragraph_spans = _join_segments(paragraphs, "index")
        return ExtractedDocument(
            text=text,
            viewer_mode="structured_text",
            mime_type=guess_mime_type(file_name, "docx"),
            meta={"paragraph_spans": paragraph_spans},
        )
    except Exception:
        fallback_text = content.decode("utf-8", errors="ignore")
        return ExtractedDocument(
            text=clean_text(fallback_text),
            viewer_mode="structured_text",
            mime_type=guess_mime_type(file_name, "docx"),
            meta={},
        )


def extract_document(content: bytes, doc_type: str, file_name: str) -> ExtractedDocument:
    if doc_type == "pdf":
        return _extract_pdf(content, file_name)
    if doc_type == "docx":
        return _extract_docx(content, file_name)

    raw = content.decode("utf-8", errors="ignore")
    if doc_type == "html":
        return ExtractedDocument(
            text=strip_html(raw),
            viewer_mode="structured_text",
            mime_type=guess_mime_type(file_name, doc_type),
            meta={},
        )

    return ExtractedDocument(
        text=clean_text(raw),
        viewer_mode="structured_text",
        mime_type=guess_mime_type(file_name, doc_type),
        meta={},
    )


def _resolve_upload_root() -> Path:
    upload_root = Path(settings.UPLOAD_DIR)
    if upload_root.is_absolute():
        return upload_root
    return Path(upload_root)


def _store_original_file(doc_id: str, file_name: str, content: bytes) -> str:
    suffix = Path(file_name).suffix or ""
    relative_path = Path("uploads") / doc_id / f"original{suffix.lower()}"
    absolute_path = _resolve_upload_root().parent / relative_path
    absolute_path.parent.mkdir(parents=True, exist_ok=True)
    absolute_path.write_bytes(content)
    return relative_path.as_posix()


def _clear_original_file_dir(doc_id: str) -> None:
    upload_dir = _resolve_upload_root().parent / "uploads" / doc_id
    if upload_dir.exists():
        shutil.rmtree(upload_dir, ignore_errors=True)


def _source_fingerprint(content: bytes) -> str:
    return hashlib.sha1(content).hexdigest()


def _jsonable_value(value: Any) -> Any:
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    return value


def import_document(
    repo,
    file_name: str,
    content: bytes,
    doc_type: Optional[str] = None,
    source_url: Optional[str] = None,
    overwrite_doc_id: Optional[str] = None,
) -> Dict[str, Any]:
    resolved_type = (doc_type or guess_doc_type(file_name)).lower()
    extracted = extract_document(content, resolved_type, file_name)
    if not extracted.text.strip():
        raise ValueError("文档解析后为空，请检查文件格式。")

    fingerprint = _source_fingerprint(content)
    overwrite_target = None
    if overwrite_doc_id:
        overwrite_target = repo.get_document(overwrite_doc_id, include_text=False)
        if overwrite_target is None:
            raise ValueError("要覆盖的文档不存在。")
        doc_id = overwrite_doc_id
        repo.clear_document_index(doc_id)
        _clear_original_file_dir(doc_id)
    else:
        duplicate = repo.find_document_by_source_fingerprint(fingerprint)
        if duplicate is not None:
            raise DuplicateDocumentError(
                {
                    "doc_id": duplicate["doc_id"],
                    "title": duplicate["title"],
                    "doc_type": duplicate["doc_type"],
                    "parse_status": duplicate["parse_status"],
                    "chunks": duplicate.get("chunks", 0),
                    "created_at": _jsonable_value(duplicate["created_at"]),
                    "original_file_name": duplicate.get("original_file_name"),
                }
            )
        doc_id = f"doc_{uuid.uuid4().hex[:12]}"
    stored_file_path = _store_original_file(doc_id, file_name, content)
    title = extract_title(file_name, extracted.text)

    meta = {
        "chunks": 0,
        "domain": "劳动法/通用",
        "source_version": 2,
        "source_fingerprint": fingerprint,
        "chunk_strategy_version": CHUNK_STRATEGY_VERSION,
        "semantic_chunking_enabled": False,
        "original_file_name": file_name,
        "mime_type": extracted.mime_type,
        "has_original_file": True,
        "viewer_mode": extracted.viewer_mode,
        **extracted.meta,
    }
    payload = {
        "doc_id": doc_id,
        "title": title,
        "doc_type": resolved_type,
        "source_url": source_url,
        "file_name": file_name,
        "file_path": stored_file_path,
        "content_text": extracted.text,
        "text": extracted.text,
        "created_at": now_iso(),
        "parse_status": "imported",
        "chunks": 0,
        "meta": meta,
    }
    repo.upsert_document(payload)
    return {
        "doc_id": doc_id,
        "title": title,
        "doc_type": resolved_type,
        "status": "imported",
        "overwritten": overwrite_target is not None,
    }
